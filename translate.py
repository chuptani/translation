import os
from docx import Document
from deep_translator import GoogleTranslator as Translator


def translate_text(text):

    text_length = len(text)
    text_parts = []
    translated_text = ""

    if text_length <= 3000:
        text_parts.append(text)
    else:
        parts = (text_length // 3000) + 2
        part_length = text_length // parts
        part_start = 0
        part_end = part_length

        for i in range(parts-1):
            for j in range(part_end, -1, -1):
                if text[j] == "।":
                    part_end = j + 1
                    break
            text_parts.append(text[part_start:part_end])
            part_start = part_end
            part_end = part_start + part_length

        text_parts.append(text[part_start:])

    for txt in text_parts:
        translated_text += Translator(source='auto', target='telugu').translate(txt)

    return translated_text


def translate_doc(path):

    doc = Document(path)
    translated_doc = Document()

    for paragraph in doc.paragraphs:
        translated_text = translate_text(paragraph.text)
        translated_doc.add_paragraph(translated_text)

    return translated_doc


if __name__ == "__main__":

    input_path = "path/to/input/folder"
    output_path = "path/to/output/folder"
    for filename in os.listdir(input_path):
        if filename.endswith(".docx"):
            input_file = os.path.join(input_path, filename)
            output_file = os.path.join(output_path, f"hindi_to_telugu_{filename}")
            translated_doc = translate_doc(input_file)
            translated_doc.save(output_file)
